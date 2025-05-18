import json
from pathlib import Path

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore

import fill_assessment


def test_fill_template(tmp_path):
    if Document is None:
        # Skip test if python-docx is not available
        return
    template_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("Name: [text_name]")
    doc.add_paragraph("[table1]")
    doc.save(template_path)

    data = {
        "name": "John",
        "table1": [["A", "B"], ["1", "2"]],
    }
    data_path = tmp_path / "data.json"
    data_path.write_text(json.dumps(data), encoding="utf-8")
    output_path = tmp_path / "out.docx"

    fill_assessment.fill_template(str(template_path), str(data_path), str(output_path))

    filled = Document(output_path)
    content = "\n".join(p.text for p in filled.paragraphs)
    assert "[text_" not in content
    assert "[table" not in content

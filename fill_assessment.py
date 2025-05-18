import argparse
import json
import logging
import re
from typing import Any, Dict, List

try:
    from docx import Document
    from docx.table import Table
except ImportError:  # pragma: no cover - library may not be installed in test env
    Document = None  # type: ignore
    Table = None  # type: ignore

try:
    from docx2pdf import convert
except ImportError:  # pragma: no cover
    convert = None  # type: ignore


logger = logging.getLogger(__name__)


def _replace_text(doc: "Document", data: Dict[str, Any]) -> None:
    pattern = re.compile(r"\[text_([^\]]+)\]")
    for para in doc.paragraphs:
        for run in para.runs:
            matches = pattern.findall(run.text)
            new_text = run.text
            for key in matches:
                value = str(data.get(key, ""))
                new_text = new_text.replace(f"[text_{key}]", value)
            if new_text != run.text:
                logger.debug("Replacing text in paragraph '%s' -> '%s'", run.text, new_text)
                run.text = new_text


def _create_table(doc: "Document", rows: List[List[Any]]) -> "Table":
    table = doc.add_table(rows=0, cols=len(rows[0]) if rows else 1)
    for row in rows:
        cells = table.add_row().cells
        for idx, val in enumerate(row):
            cells[idx].text = str(val)
    return table


def _replace_tables(doc: "Document", data: Dict[str, Any]) -> None:
    pattern = re.compile(r"\[table(\d+)\]")
    for para in list(doc.paragraphs):  # copy since we modify structure
        match = pattern.search(para.text)
        if not match:
            continue
        table_key = f"table{match.group(1)}"
        rows = data.get(table_key)
        if not isinstance(rows, list):
            logger.warning("No data for %s", table_key)
            para.text = para.text.replace(match.group(0), "")
            continue
        logger.debug("Inserting table for %s", table_key)
        table = _create_table(doc, rows)
        tbl_element = table._element
        tbl_element.getparent().remove(tbl_element)
        para._p.addnext(tbl_element)
        # remove placeholder paragraph if it's just the tag
        if para.text.strip() == match.group(0):
            para._element.getparent().remove(para._element)
        else:
            para.text = para.text.replace(match.group(0), "")


def fill_template(input_path: str, data_path: str, output_path: str, preview: bool = False) -> None:
    if Document is None:
        raise RuntimeError("python-docx is required to run this script")
    logger.info("Loading template %s", input_path)
    try:
        doc = Document(input_path)
    except Exception as exc:  # pragma: no cover - docx may raise various errors
        logger.error("Failed to load template: %s", exc)
        raise

    try:
        with open(data_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except Exception as exc:
        logger.error("Failed to read JSON data: %s", exc)
        raise

    _replace_text(doc, data)
    _replace_tables(doc, data)

    logger.info("Saving filled document to %s", output_path)
    try:
        doc.save(output_path)
    except Exception as exc:
        logger.error("Failed to save document: %s", exc)
        raise

    if preview:
        if convert is None:
            logger.error("docx2pdf package not available, cannot generate preview")
        else:
            try:
                logger.info("Generating PDF preview")
                convert(output_path)
            except Exception as exc:
                logger.error("Failed to generate PDF preview: %s", exc)


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill assessment template")
    parser.add_argument("--input", required=True, help="Path to the docx template")
    parser.add_argument("--data", required=True, help="Path to JSON file with values")
    parser.add_argument("--output", required=True, help="Output docx path")
    parser.add_argument("--preview", action="store_true", help="Generate PDF preview")
    parser.add_argument("--log", default="INFO", help="Logging level")
    args = parser.parse_args()

    logging.basicConfig(level=getattr(logging, args.log.upper(), logging.INFO))
    try:
        fill_template(args.input, args.data, args.output, args.preview)
    except Exception as exc:  # pragma: no cover - entry point wrapper
        logger.error("Error filling template: %s", exc)
        raise


if __name__ == "__main__":
    main()
